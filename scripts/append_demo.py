"""
分析 docx body 的 XML 结构 + Markdown→docx 追加原型。
展示如何在封面后追加内容，为正式写入做准备。
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as ET


DOCS_DIR = Path(__file__).parent.parent / "docs"
OUTPUT_DIR = Path(__file__).parent.parent / "output"


def print_body_structure(doc):
    """打印 docx body 的 XML 子元素结构，理解元素排列方式。"""
    body = doc.element.body
    print("=" * 60)
    print("Body 子元素 (XML tag 顺序):")
    print("=" * 60)

    para_idx = 0
    table_idx = 0
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        text_preview = ""
        if tag == "p":
            # 找对应段落文本
            if para_idx < len(doc.paragraphs):
                text_preview = doc.paragraphs[para_idx].text[:40]
                text_preview = text_preview if text_preview.strip() else "(空)"
            print(f"  [{para_idx:2d}] <w:p>        「{text_preview}」")
            para_idx += 1
        elif tag == "tbl":
            rows = len(doc.tables[table_idx].rows) if table_idx < len(doc.tables) else "?"
            cols = len(doc.tables[table_idx].columns) if table_idx < len(doc.tables) else "?"
            print(f"        <w:tbl>      表格{table_idx} ({rows}r × {cols}c)")
            table_idx += 1
        elif tag == "sectPr":
            print(f"        <w:sectPr>   (节属性 - 页面设置)")
        else:
            print(f"        <w:{tag}>")


def append_paragraph(doc, text, style="Normal", font_name="宋体", font_size=Pt(12),
                     bold=False, alignment=None, space_after=Pt(6)):
    """在文档末尾追加一个段落，返回段落对象。"""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    # 设置中文字体 (w:rFonts 的 eastAsia 属性)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p


def append_heading(doc, text, level=2):
    """追加标题段落 (用 Normal 样式模拟 heading，因为模板没有 heading 样式)。"""
    size_map = {1: Pt(18), 2: Pt(15), 3: Pt(13)}
    return append_paragraph(doc, text, font_name="黑体",
                           font_size=size_map.get(level, Pt(12)),
                           bold=True, space_after=Pt(12))


def append_code_block(doc, code_text):
    """追加代码块（灰色背景 + 等宽字体）。"""
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph()
        # 段前段后间距最小
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)

        # 灰色底纹 (shd)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)

        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), "宋体")  # 中文回退

    # 代码块后加一个空行
    append_paragraph(doc, "", font_size=Pt(6), space_after=Pt(6))


def append_table_from_md(doc, headers, rows):
    """追加表格（从 Markdown 表格转换）。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h.strip()
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val.strip()
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # 表后空行
    append_paragraph(doc, "", font_size=Pt(6), space_after=Pt(6))
    return table


def append_bullet(doc, text):
    """追加一个列表项。"""
    p = doc.add_paragraph()
    run = p.add_run("• " + text)
    run.font.name = "宋体"
    run.font.size = Pt(12)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), "宋体")
    rPr.insert(0, rFonts)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    return p


def demo_append():
    """演示：打开模板，在末尾追加一段测试内容，保存到 output。"""
    src = DOCS_DIR / "2026语法分析实习报告封面 - 软工23 (1).docx"
    out = OUTPUT_DIR / "_demo_appended.docx"

    doc = Document(str(src))

    # ── 先分析 body 结构 ──
    print_body_structure(doc)
    print(f"\n最后段落: P{len(doc.paragraphs)-1} = 「{doc.paragraphs[-1].text[:50]}」")
    print(f"最后表格: 共{len(doc.tables)}个")

    # sectPr 是 body 的最后一个子元素，add_* 会自动在 sectPr 之前插入
    # 所以直接 add_paragraph / add_table 就行！

    # ── 追加演示内容 ──
    append_heading(doc, "第一部分  语言语法规则（自然语言描述）", level=1)
    append_paragraph(doc, "这是一段正文内容，展示如何在封面后追加普通段落。"
                     "采用宋体 12pt，段后间距 6pt。")

    append_heading(doc, "1.1 语法变量表", level=2)
    append_paragraph(doc, "AST 各节点类型如下：")

    append_heading(doc, "顶层结构", level=3)
    append_table_from_md(doc,
        headers=["类型", "字段", "说明"],
        rows=[
            ["Program", "stms", "语句序列的向量"],
        ])

    append_heading(doc, "语句（Statement）", level=3)
    append_table_from_md(doc,
        headers=["类型", "字段", "说明"],
        rows=[
            ["DeclStatement", "name: String, init: Expr", "声明语句"],
            ["AssignStatement", "name: String, value: Expr", "赋值语句"],
            ["IfStatement", "cond: Expr, then, else?", "if 分支"],
            ["WhileStatement", "cond: Expr, body: Block", "while 循环"],
        ])

    append_heading(doc, "测试代码", level=2)
    append_code_block(doc, 'i32 a = 1;\nif (a == 1) {\n    a = a + 1;\n}')

    append_heading(doc, "核心设计决策", level=2)
    append_bullet(doc, "BoolExpr 限定为关系表达式")
    append_bullet(doc, "Block 强制使用花括号")
    append_bullet(doc, "运算符优先级通过文法编码")

    doc.save(str(out))
    print(f"\n✅ 演示文件已保存: {out}")


if __name__ == "__main__":
    demo_append()
