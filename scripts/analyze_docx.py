"""
深度分析 docs 目录下的 docx 报告模板结构，为写入做准备。
用法: uv run python scripts/analyze_docx.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def emu_to_cm(emu):
    if emu is None: return "N/A"
    return f"{emu/360000:.2f} cm"


def emu_to_pt(emu):
    if emu is None: return "N/A"
    return f"{emu/12700:.1f} pt"


def alignment_str(align):
    m = {None: "默认", WD_ALIGN_PARAGRAPH.LEFT: "左", WD_ALIGN_PARAGRAPH.CENTER: "中",
         WD_ALIGN_PARAGRAPH.RIGHT: "右", WD_ALIGN_PARAGRAPH.JUSTIFY: "两端", WD_ALIGN_PARAGRAPH.DISTRIBUTE: "分散"}
    return m.get(align, str(align))


def main():
    docs_dir = Path(__file__).parent.parent / "docs"
    docx_files = sorted(docs_dir.glob("*.docx"))
    if not docx_files:
        print("没有 .docx 文件")
        return

    fp = docx_files[0]
    doc = Document(str(fp))
    print(f"文件: {fp.name}\n")

    # ========== 1. 页面设置 ==========
    sec = doc.sections[0]
    print("=" * 60)
    print("页面设置")
    print(f"  纸张: {emu_to_cm(sec.page_width)} × {emu_to_cm(sec.page_height)}")
    print(f"  边距: 上{emu_to_cm(sec.top_margin)} 下{emu_to_cm(sec.bottom_margin)} 左{emu_to_cm(sec.left_margin)} 右{emu_to_cm(sec.right_margin)}")
    print(f"  页眉距离: {emu_to_cm(sec.header_distance)}  页脚距离: {emu_to_cm(sec.footer_distance)}")

    # ========== 2. 段落逐一分析 (含 oxml 关键信息) ==========
    print(f"\n{'='*60}")
    print(f"段落结构 (共 {len(doc.paragraphs)} 段)")
    print("=" * 60)

    for i, p in enumerate(doc.paragraphs):
        text = p.text
        style = p.style.name
        align = alignment_str(p.alignment)
        pf = p.paragraph_format
        ls = pf.line_spacing
        if ls:
            ls = f"{ls:.1f}" if isinstance(ls, float) else str(ls)

        marker = ""
        if not text.strip():
            marker = " [空]"
        print(f"\nP{i} | style={style} align={align} line_spacing={ls}{marker}")
        if text.strip():
            print(f"  全文: 「{text}」")

        # 每个 run 详情
        if p.runs:
            for j, r in enumerate(p.runs):
                f = r.font
                parts = []
                if f.name: parts.append(f"字体={f.name}")
                if f.size: parts.append(f"大小={emu_to_pt(f.size)}")
                if f.bold: parts.append("粗")
                if f.italic: parts.append("斜")
                if f.underline: parts.append("下划线")
                if f.color and f.color.rgb: parts.append(f"色=#{f.color.rgb}")
                rt = r.text.replace("\n","\\n")
                print(f"  R{j}: 「{rt[:80]}」 {' | '.join(parts)}")

    # ========== 3. 表格分析 (含列宽) ==========
    print(f"\n{'='*60}")
    print(f"表格结构 (共 {len(doc.tables)} 个)")
    print("=" * 60)

    for ti, t in enumerate(doc.tables):
        print(f"\n表格 {ti}: {len(t.rows)}行 × {len(t.columns)}列, style={t.style.name}")

        # 列宽
        widths = []
        for ci in range(len(t.columns)):
            w = t.columns[ci].width
            widths.append(w)
        print(f"  列宽: {', '.join(emu_to_cm(w) for w in widths)}")

        # 单元格详细
        for ri, row in enumerate(t.rows):
            print(f"  行{ri}:")
            for ci, cell in enumerate(row.cells):
                ct = cell.text.replace("\n", "\\n")[:40]
                # 合并单元格检测
                tcPr = cell._tc.find(qn('w:tcPr'))
                grid_span = vmerge = hmerge = None
                if tcPr is not None:
                    gs = tcPr.find(qn('w:gridSpan'))
                    if gs is not None:
                        grid_span = int(gs.get(qn('w:val')))
                    vm = tcPr.find(qn('w:vMerge'))
                    if vm is not None:
                        vmerge = vm.get(qn('w:val')) or 'continue'
                    hm = tcPr.find(qn('w:hMerge'))
                    if hm is not None:
                        hmerge = hm.get(qn('w:val')) or 'continue'

                # 单元格内段落对齐
                aligns = [alignment_str(pp.alignment) for pp in cell.paragraphs]
                tags = []
                if grid_span: tags.append(f"跨{grid_span}列")
                if vmerge: tags.append(f"垂合:{vmerge}")
                if hmerge: tags.append(f"水合:{hmerge}")
                tag_str = f" [{', '.join(tags)}]" if tags else ""
                print(f"    [{ri},{ci}] align={aligns} 「{ct}」{tag_str}")

    # ========== 4. 写入关键点总结 ==========
    print(f"\n{'='*60}")
    print("写入定位总结")
    print("=" * 60)

    fields = [
        ("编号", "P0 R1", "段落0 下划线Run (14pt 下划线) → 写入编号字符串"),
    ]

    # 表格字段
    grade_cols = {1:"一",2:"二",3:"三",4:"四",5:"五",6:"六",7:"七",8:"八",9:"九",10:"十"}
    for ci in range(1, 11):
        fields.append((f"实习{grade_cols[ci]}成绩", f"表0 [1,{ci}]", f"表格第1行第{ci}列 → 写入成绩"))
    fields.append(("总评", "表0 [1,11]", "表格第1行第11列 → 写入总评"))
    fields.append(("教师签名", "表0 [1,12]", "表格第1行第12列 → 写入签名"))

    for label, loc, desc in fields:
        print(f"  {loc:15s} {label}")

    print()
    # 段落字段 (已填内容若要替换的位置)
    text_fields = [
        ("编号/班序号", "P9",  "编    号：  班序号 CP      → 替换 CP 为实际班序号"),
        ("实习题目",     "P10", "实习题目：               → 填入下划线区域"),
        ("专业(班)",     "P11", "专业（班）：              → 填入下划线区域"),
        ("学生学号",     "P12", "学生学号：               → 填入下划线区域"),
        ("学生姓名",     "P13", "学生姓名：               → 填入下划线区域"),
        ("任课教师",     "P14", "任课教师： 杜卓敏         → 修改教师名"),
        ("日期(年)",     "P20", "Run1=６ → 修改年份最后一位; Run3=空格 → 月; Run5=空格 → 日"),
    ]
    for label, loc, desc in text_fields:
        print(f"  {loc:6s} {label:<12s} {desc}")

    print(f"\n  目录页码: P25-P29 最后一个 Run 的文本")
    print(f"    P25 含页号 '1', P26-P29 页号为空, 需回填")


if __name__ == "__main__":
    main()
