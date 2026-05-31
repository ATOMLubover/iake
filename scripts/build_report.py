"""
构建实习报告：封面填充 + 规范排版正文。

用法:
    uv run python scripts/build_report.py
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT = Path(__file__).resolve().parent.parent
TEMPLATE = PROJECT / "docs" / "2026语法分析实习报告封面 - 软工23 (1).docx"
OUTPUT = PROJECT / "output" / "语法分析实习报告.docx"


# ═══════════════════════════════════════════════════════════════
# 封面配置
# ═══════════════════════════════════════════════════════════════

@dataclass
class CoverConfig:
    top_number: str = ""
    class_number: str = "软工23-CP"
    topic: str = "Mini 语言语法分析器"
    major: str = "软件工程"
    student_id: str = ""
    student_name: str = ""
    teacher: str = "杜卓敏"
    year: int | None = None
    month: int | None = None
    day: int | None = None
    toc_pages: tuple[int, ...] = (1, 0, 0, 0, 0)


# ═══════════════════════════════════════════════════════════════
# 报告内容定义（分段编排，数据来自 report-parser.md）
# ═══════════════════════════════════════════════════════════════

# 表格数据
TBL_TOP_LEVEL = (
    ["类型", "字段", "说明"],
    [
        ["Program", "stms", "语句序列的向量"],
    ])

TBL_STATEMENT = (
    ["类型", "字段", "说明"],
    [
        ["DeclStatement",   "name: String, init: Expr",   "声明语句"],
        ["AssignStatement", "name: String, value: Expr",  "赋值语句"],
        ["IfStatement",     "cond: Expr, then, else?",    "if 分支"],
        ["WhileStatement",  "cond: Expr, body: Block",    "while 循环"],
    ])

TBL_EXPRESSION = (
    ["类型", "字段", "说明"],
    [
        ["Identifier", "name: String",                     "标识符引用"],
        ["Integer",    "value: i64",                       "整数字面量"],
        ["Binary",     "oper: BinaryOperator, left, right","二元运算"],
    ])

TBL_BINARY_OP = (
    ["运算符", "Token", "说明"],
    [
        ["Equal", "==", "相等比较"],
        ["Less",  "<",  "小于比较"],
        ["Add",   "+",  "加法"],
        ["Mul",   "*",  "乘法"],
    ])

TBL_BLOCK = (
    ["类型", "字段", "说明"],
    [
        ["Block", "stms", "{ StmList }"],
    ])

TBL_TERMINALS = (
    ["记号", "含义"],
    [
        ["i32",   "类型关键字"],
        ["if",    "条件分支"],
        ["else",  "否则分支"],
        ["while", "循环"],
        ["id",    "标识符"],
        ["num",   "非负整数"],
        ["=",     "赋值"],
        ["==",    "相等比较"],
        ["<",     "小于比较"],
        ["+",     "加法"],
        ["*",     "乘法"],
        ["( )",   "圆括号"],
        ["{ }",   "语句块"],
        [";",     "语句结束"],
        ["$",     "输入结束符"],
    ])

TBL_PRODUCTIONS = (
    ["编号", "产生式"],
    [
        [ "(1)",  "Program → StmList"],
        [ "(2)",  "StmList → Stm StmList"],
        [ "(3)",  "StmList → ε"],
        [ "(4)",  "Stm → DeclStm"],
        [ "(5)",  "Stm → AssignStm"],
        [ "(6)",  "Stm → IfStm"],
        [ "(7)",  "Stm → WhileStm"],
        [ "(8)",  "DeclStm → i32 id = ArithExpr ;"],
        [ "(9)",  "AssignStm → id = ArithExpr ;"],
        ["(10)", "IfStm → if ( BoolExpr ) Block ElsePart"],
        ["(11)", "ElsePart → else Block"],
        ["(12)", "ElsePart → ε"],
        ["(13)", "WhileStm → while ( BoolExpr ) Block"],
        ["(14)", "Block → { StmList }"],
        ["(15)", "BoolExpr → ArithExpr RelOp ArithExpr"],
        ["(16)", "RelOp → =="],
        ["(17)", "RelOp → <"],
        ["(18)", "ArithExpr → Term ArithTail"],
        ["(19)", "ArithTail → + Term ArithTail"],
        ["(20)", "ArithTail → ε"],
        ["(21)", "Term → Factor TermTail"],
        ["(22)", "TermTail → * Factor TermTail"],
        ["(23)", "TermTail → ε"],
        ["(24)", "Factor → id"],
        ["(25)", "Factor → num"],
        ["(26)", "Factor → ( ArithExpr )"],
    ])

TBL_ERRORS = (
    ["错误类型", "触发位置", "含义"],
    [
        ["ParserError::Lex", "fill_ahead()", "Lexer 产生词法错误，Parser 透传并附带 cursor"],
        ["ParserError::UnexpectedToken", "各 expect_* / parse_*", "期望某类 token，但遇到其他类型的 token"],
        ["ParserError::UnexpectedEndOfInput", "各 expect_* / parse_*", "期望 token，但输入流已耗尽（EOF）"],
    ])

TBL_UNIT_TESTS = (
    ["测试维度", "测试内容", "示例"],
    [
        ["声明语句",      "i32 + 标识符 + = + 表达式 + ;",                                        'i32 a = 1;'],
        ["运算符优先级",  "* 优先于 + 的结合次序验证",                                             'a = b + c * 2; → Add(b, Mul(c, 2))'],
        ["嵌套控制流",    "if / while / else 完整嵌套",                                            'if (a == 1) { while (b < 10) { ... } } else { ... }'],
        ["缺失分号",      "遗漏 ; 的错误检测",                                                     'a = 1 → error'],
        ["缺失右括号",    "if 条件缺少 )",                                                         'if (a == 1 { ... } → error'],
        ["缺失右花括号",  "while body 缺少 }",                                                     'while (a < 1) { a = a + 1; → error'],
        ["缺失左括号",    "if 后没有 (",                                                           'if a == 1) { ... } → error'],
        ["空赋值值",      "= 后无表达式",                                                          'a = ; → error'],
        ["BoolExpr 缺关系运算符", "while 条件中缺 == 或 <",                                         'while (a) { ... } → error'],
        ["词法错误穿透",  "非法字符被 Lexer 捕获并由 Parser 透传",                                  'a = @; → InvalidChar(\'@\')'],
    ])

TBL_INTEGRATION = (
    ["用例", "输入描述", "预期输出", "验证要点"],
    [
        ["input-1", "含注释、声明、if-else、while 嵌套的完整程序",
         "完整前序语法树", "复杂控制流 + 注释跳过 + 嵌套 AST 结构正确"],
        ["input-2", "含非法字符 @ 的程序",
         "位置为 (line 1, col 8) 的 InvalidChar('@')", "词法错误穿透 Parser 并准确定位"],
        ["input-3", "带前导零的整数 0123",
         "位置为 (line 0, col 9) 的 UnexpectedChar('1')", "Lexer 前导零拒绝策略 → Parser 错误透传"],
        ["input-4", "第一行遗漏分号",
         "(line 1, col 3): expected ;, found Keyword(i32)", "语义错误定位"],
        ["input-5", "遗漏右花括号直到 EOF",
         "(line 2, col 0): expected }, found EOF", "EOF 错误检测"],
    ])


# ═══════════════════════════════════════════════════════════════
# Docx 构建工具
# ═══════════════════════════════════════════════════════════════

class DocBuilder:
    """流式构建 docx 正文，自动管理标题编号和格式。"""

    def __init__(self, doc: Document):
        self.doc = doc
        self.counters: dict[int, int] = {}  # level → counter

    # ── 字体工具 ──

    def _set_font(self, run, name: str, size, bold: bool = False):
        run.font.name = name
        run.font.size = size
        run.bold = bold
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), name)

    def _run(self, p, text: str, name: str, size, bold: bool = False):
        """在段落中加一个 run。"""
        r = p.add_run(text)
        self._set_font(r, name, size, bold)
        return r

    # ── 标题编号 ──

    def _advance_counter(self, level: int) -> str:
        """推进计数器，返回如 '1.2.3' 的编号字符串。"""
        # 重置更深层计数器
        for lv in list(self.counters.keys()):
            if lv > level:
                del self.counters[lv]
        self.counters[level] = self.counters.get(level, 0) + 1
        # 组装编号
        parts = [str(self.counters[lv]) for lv in sorted(self.counters)]
        return ".".join(parts)

    # ── 段落级元素 ──

    def heading1(self, text: str):
        """第一部分 标题（18pt 黑体，段前分页）。"""
        num = self._advance_counter(1)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        # 分页（第一个 h1 除外）
        self._run(p, f"{num}  ", "黑体", Pt(18), bold=True)
        self._run(p, text, "黑体", Pt(18), bold=True)

    def heading2(self, text: str):
        """二级标题（15pt 黑体）。"""
        num = self._advance_counter(2)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        self._run(p, f"{num}  ", "黑体", Pt(15), bold=True)
        self._run(p, text, "黑体", Pt(15), bold=True)

    def heading3(self, text: str):
        """三级标题（13pt 黑体）。"""
        num = self._advance_counter(3)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        self._run(p, f"{num}  ", "黑体", Pt(13), bold=True)
        self._run(p, text, "黑体", Pt(13), bold=True)

    def para(self, text: str, indent: bool = False):
        """正文段落（12pt 宋体）。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
        self._run(p, text, "宋体", Pt(12))

    def code(self, text: str):
        """代码块（Courier New 8.5pt，灰底）。"""
        lines = text.strip().split("\n")
        for line in lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(14)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F0F0")
            pPr.append(shd)
            self._run(p, line if line else " ", "Courier New", Pt(8.5))

    def table(self, headers: list[str], rows: list[list[str]]):
        """表格（Table Grid，表头灰底加粗）。"""
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ci, h in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(h)
            self._set_font(r, "宋体", Pt(10), bold=True)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "E8E8E8")
            tcPr.append(shd)

        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = tbl.rows[ri + 1].cells[ci]
                cell.text = ""
                r = cell.paragraphs[0].add_run(val)
                self._set_font(r, "宋体", Pt(10))

        # 表后留空
        self.para("")

    def page_break(self):
        """分页符。"""
        p = self.doc.add_paragraph()
        run = p.add_run()
        run._r.append(OxmlElement("w:br"))
        run._r[-1].set(qn("w:type"), "page")

    def bullet(self, text: str):
        """无序列表项。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1)
        self._run(p, "• " + text, "宋体", Pt(12))

    def numbered(self, items: list[str]):
        """有序列表。"""
        for i, text in enumerate(items, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(1)
            self._run(p, f"{i}. {text}", "宋体", Pt(12))
        self.para("")


# ═══════════════════════════════════════════════════════════════
# 封面填充
# ═══════════════════════════════════════════════════════════════

def fill_cover(doc: Document, cfg: CoverConfig):
    def to_fullwidth(n: int) -> str:
        return "".join(chr(ord("０") + int(d)) for d in str(n))

    paras = doc.paragraphs

    def _fill(p_idx: int, value: str):
        p = paras[p_idx]
        runs = p.runs
        if len(runs) < 2:
            return
        ow = sum(len(r.text) for r in runs[1:])
        padded = (value + " " * max(0, ow - len(value))) if value else " " * ow
        runs[1].text = padded
        for r in runs[2:]:
            r.text = ""

    if cfg.top_number:
        _fill(0, cfg.top_number)
    _fill(9, cfg.class_number)
    _fill(10, cfg.topic)
    _fill(11, cfg.major)
    _fill(12, cfg.student_id)
    _fill(13, cfg.student_name)
    _fill(14, cfg.teacher)

    now = date.today()
    y = cfg.year or now.year
    m = cfg.month or now.month
    d = cfg.day or now.day
    full = f"{to_fullwidth(y)}  年  {to_fullwidth(m)}  月  {to_fullwidth(d)}  日"
    runs = paras[20].runs
    runs[0].text = full
    for r in runs[1:]:
        r.text = ""

    for i, page in enumerate(cfg.toc_pages):
        if page > 0 and (25 + i) < len(paras):
            p = paras[25 + i]
            if p.runs:
                last = p.runs[-1]
                t = last.text
                j = len(t) - 1
                while j >= 0 and t[j].isdigit():
                    j -= 1
                if j + 1 < len(t):
                    last.text = t[:j + 1] + str(page)


# ═══════════════════════════════════════════════════════════════
# 报告正文
# ═══════════════════════════════════════════════════════════════

def build_body(b: DocBuilder):
    """按规范结构编写报告正文。"""

    # ═══════════════════════════════════════════════════════════
    # 第一部分：语言语法规则
    # ═══════════════════════════════════════════════════════════
    b.heading1("语言语法规则（自然语言描述）")

    b.para(
        "本实验设计了一门名为 Mini 的微型命令式语言。"
        "该语言支持整型变量的声明与赋值、算术与关系表达式求值、"
        "if-else 条件分支以及 while 循环控制流，"
        "语句以分号结束，语句块由花括号界定。"
        "语言采用词法分析和语法分析两阶段处理，"
        "空白符和 # 单行注释在词法分析阶段被消去，不进入语法分析阶段。",
        indent=True,
    )

    b.heading2("语法变量表")
    b.para("抽象语法树（AST）各节点类型定义如下。", indent=True)

    b.heading3("顶层结构")
    b.table(*TBL_TOP_LEVEL)

    b.heading3("语句（Statement）")
    b.table(*TBL_STATEMENT)

    b.heading3("表达式（Expression）")
    b.table(*TBL_EXPRESSION)

    b.heading3("二元运算符（BinaryOperator）")
    b.table(*TBL_BINARY_OP)

    b.heading3("辅助结构")
    b.table(*TBL_BLOCK)

    # ═══════════════════════════════════════════════════════════
    # 第二部分：文法定义
    # ═══════════════════════════════════════════════════════════
    b.page_break()
    b.heading1("文法定义")

    b.heading2("LL(1) 文法概述")
    b.para(
        "开始符号为 Program，文法采用递归下降可解析的 LL(1) 形式。"
        "每个 (非终结符, 向前看) 组合至多对应一条产生式，满足 LL(1) 要求。",
        indent=True,
    )

    b.heading2("终结符约定")
    b.para(
        "以下终结符来自 Lexer 产出的 Token。"
        "空白符和 # 单行注释在词法分析阶段已被消去，不进入语法分析。",
        indent=True,
    )
    b.table(*TBL_TERMINALS)

    b.heading2("产生式表")
    b.para("完整产生式如下表所示，共计 26 条产生式。", indent=True)
    b.table(*TBL_PRODUCTIONS)

    b.heading2("核心设计决策")
    b.para("文法设计过程中做出以下关键决策：", indent=True)

    b.numbered([
        "BoolExpr 限定为关系表达式：BoolExpr → ArithExpr RelOp ArithExpr，"
        "不涉及短路逻辑，由此简化了整个文法。",

        "Block 强制使用花括号：Block → { StmList }，"
        "因此不存在悬空 else 问题——ElsePart 的 else 总是匹配最近的 if。",

        "运算符优先级通过文法编码：文法将 ArithExpr 拆分为 Term（含 *）"
        "和 ArithTail（含 +），天然编码了 * 高于 + 的优先级。",
    ])

    # ═══════════════════════════════════════════════════════════
    # 第三部分：语法分析算法
    # ═══════════════════════════════════════════════════════════
    b.page_break()
    b.heading1("语法分析算法")

    b.heading2("Parser 总体设计")
    b.para(
        "Parser 的核心是将 Lexer 产出的 Token 流转化为 AST。"
        "与 Lexer 流式 next_token() 的接口配合，"
        "Parser 在工程上采用递归下降 + 单 token 前瞻的方案，"
        "这是最自然高效的选择。",
        indent=True,
    )

    b.para("核心接口定义如下：")
    b.code("""pub struct Parser<I: Input> {
    lexer: Lexer<I>,
    ahead: Option<Token>,  // 单 token 前瞻缓冲区
    cursor: Cursor,
}

impl<I: Input> Parser<I> {
    pub fn new(input: I) -> Self { /* ... */ }
    pub fn parse_program(&mut self) -> ParserResult<Program> { /* ... */ }
}""")

    b.heading2("递归下降函数编排")
    b.para(
        "每个非终结符对应一个解析函数，向前看一个 token 决定分发方向。"
        "以下逐一列出全部解析函数的伪代码。",
        indent=True,
    )

    b.heading3("入口：parse_program")
    b.code("""parse_program():
    stms = parse_stm_list()
    if peek() is not None:
        error("end of input")
    return Program { stms }""")

    b.heading3("语句层派发：parse_stm")
    b.code("""parse_stm():
    match peek():
        Some(Keyword(I32))  → parse_decl_stm()
        Some(Identifier(_)) → parse_assign_stm()
        Some(Keyword(If))   → parse_if_stm()
        Some(Keyword(While))→ parse_while_stm()
        _                   → error("start of statement")""")

    b.heading3("声明语句：parse_decl_stm")
    b.code("""parse_decl_stm():
    expect_keyword(I32)
    name = expect_identifier()
    expect_operator(Assign)
    init = parse_arith_expr()
    expect_punctuation(Semicolon)
    return DeclStatement { name, init }""")

    b.heading3("赋值语句：parse_assign_stm")
    b.code("""parse_assign_stm():
    name = expect_identifier()
    expect_operator(Assign)
    value = parse_arith_expr()
    expect_punctuation(Semicolon)
    return AssignStatement { name, value }""")

    b.heading3("if 分支：parse_if_stm 与 parse_else_part")
    b.code("""parse_if_stm():
    expect_keyword(If)
    expect_punctuation(ParenLeft)
    cond = parse_bool_expr()
    expect_punctuation(ParenRight)
    then_block = parse_block()
    else_block = parse_else_part()
    return IfStatement { cond, then_block, else_block }

parse_else_part():
    match peek():
        Some(Keyword(Else)):
            expect_keyword(Else)
            block = parse_block()
            return Some(block)
        _:
            return None     // ε：无 else 子句""")

    b.heading3("while 循环：parse_while_stm")
    b.code("""parse_while_stm():
    expect_keyword(While)
    expect_punctuation(ParenLeft)
    cond = parse_bool_expr()
    expect_punctuation(ParenRight)
    body = parse_block()
    return WhileStatement { cond, body }""")

    b.heading3("语句块：parse_block")
    b.code("""parse_block():
    expect_punctuation(BraceLeft)
    stms = parse_stm_list()
    expect_punctuation(BraceRight)
    return Block { stms }""")

    b.heading3("布尔表达式：parse_bool_expr")
    b.code("""parse_bool_expr():
    left = parse_arith_expr()
    oper = parse_rel_oper()
    right = parse_arith_expr()
    return Binary { oper, left, right }""")

    b.heading3("算术表达式（带优先级）")
    b.para(
        "运算符优先级通过两层递归实现：parse_arith_expr 调用 parse_term "
        "后循环消费 +，而 parse_term 调用 parse_factor 后循环消费 *。"
        "这保证了 * 先于 + 结合，例如 a + b * c 解析为 Add(a, Mul(b, c))。",
        indent=True,
    )
    b.code("""parse_arith_expr():
    left = parse_term()
    while peek() == Operator(Add):
        expect_operator(Add)
        right = parse_term()
        left = Binary { oper: Add, left, right }
    return left

parse_term():
    left = parse_factor()
    while peek() == Operator(Mul):
        expect_operator(Mul)
        right = parse_factor()
        left = Binary { oper: Mul, left, right }
    return left

parse_factor():
    match take_token():
        Some(Identifier(name)) → Identifier { name }
        Some(Integer(value))   → Integer { value }
        Some(Punctuation(ParenLeft)):
            expr = parse_arith_expr()
            expect_punctuation(ParenRight)
            return expr
        Some(found) → error("identifier, integer, or (")
        None        → error("identifier, integer, or (")""")

    b.heading2("Token 前瞻机制")
    b.para(
        "Parser 维护一个 ahead: Option<Token> 单 token 前瞻缓冲区。"
        "采用惰性填充策略：仅当 ahead 为空时才从 Lexer 拉取。"
        "每个 expect_* 函数调用 take_token() 消费一个 token 并验证类型，"
        "不匹配时立即报告错误并附带光标位置。",
        indent=True,
    )

    b.code("""peek_token():
    fill_ahead()
    return ahead.as_ref()

take_token():
    fill_ahead()
    cursor = self.cursor
    return (ahead.take(), cursor)

fill_ahead():
    if ahead.is_some(): return
    match lexer.next_token():
        Ok(Some(tok)):
            cursor = lexer.cursor()
            ahead = Some(tok)
        Ok(None):
            cursor = lexer.cursor()
            ahead = None
        Err(err) → ParserError::Lex { err, cursor }""")

    b.heading2("AST 前序输出")
    b.para(
        "AST 提供了 preorder_string() 方法，将语法树以前序树形格式输出。"
        "每个节点调用 push_tree_line 写入一行带缩进和连接符（├── 或 └──）"
        "的文本。非终节点递归展开其子节点，通过 prefix 参数传递缩进延续字符"
        "（│   或     ）。",
        indent=True,
    )

    b.code("""Program
├── DeclStatement(a)
│   └── Integer(1)
└── IfStatement
    ├── Condition
    │   └── Equal
    │       ├── Identifier(a)
    │       └── Integer(1)
    ├── ThenBlock
    │   └── ...
    └── ElseBlock
        └── ...""")

    # ═══════════════════════════════════════════════════════════
    # 第四部分：出错处理
    # ═══════════════════════════════════════════════════════════
    b.page_break()
    b.heading1("出错处理出口")

    b.heading2("错误类型体系")
    b.para(
        "Parser 定义了三种错误类型，覆盖词法透传、非预期 Token 和 EOF 三种场景。"
        "错误输出均包含 (line, col) 光标位置信息，便于定位问题。",
        indent=True,
    )
    b.table(*TBL_ERRORS)

    b.para("错误输出格式示例如下：")
    b.bullet("(line 1, col 3): expected ;, found Keyword(i32) —— 遗漏分号")
    b.bullet("(line 2, col 0): expected }, found EOF —— 遗漏右花括号")
    b.bullet("(line 1, col 8): InvalidChar('@') —— 词法阶段捕获的非法字符")

    # ═══════════════════════════════════════════════════════════
    # 第五部分：测试计划
    # ═══════════════════════════════════════════════════════════
    b.page_break()
    b.heading1("测试计划（报告）")

    b.para(
        "测试策略为自底向上：先测单个语法结构，再测组合与控制流嵌套，"
        "最后测错误路径。",
        indent=True,
    )

    b.heading2("单元测试")
    b.para(
        "单元测试覆盖声明语句、运算符优先级、嵌套控制流等正常路径，"
        "以及缺失分号、缺失括号、空赋值、词法错误穿透等异常路径。",
        indent=True,
    )
    b.table(*TBL_UNIT_TESTS)

    b.heading2("集成测试")
    b.para(
        "集成测试按 tests/example/ 目录下的 5 对 input-N.txt / output-N.txt "
        "文件进行验证，覆盖从词法到语法分析的完整流水线。",
        indent=True,
    )
    b.table(*TBL_INTEGRATION)


# ═══════════════════════════════════════════════════════════════
# 入口
# ═══════════════════════════════════════════════════════════════

def build():
    if not TEMPLATE.exists():
        print(f"❌ 模板不存在: {TEMPLATE}")
        return

    doc = Document(str(TEMPLATE))

    # 填封面
    cfg = CoverConfig()
    fill_cover(doc, cfg)
    print("封面填充完成")

    # 写正文
    b = DocBuilder(doc)
    build_body(b)
    print("正文编排完成")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"✅ {OUTPUT}")


if __name__ == "__main__":
    build()
